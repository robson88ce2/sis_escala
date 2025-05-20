import sqlite3
import json
import os
from datetime import datetime, timedelta
import pandas as pd
import tempfile
from docx import Document
from docx2pdf import convert
from docx.shared import Inches
import subprocess
import platform
from docx.oxml import parse_xml
from docx.oxml.ns import qn
from docx.oxml.table import CT_Tbl

# --- Banco ---
DB_PATH = "escala.db"

def conectar():
    conn = sqlite3.connect(DB_PATH, detect_types=sqlite3.PARSE_DECLTYPES | sqlite3.PARSE_COLNAMES)
    conn.row_factory = sqlite3.Row
    return conn

def criar_tabelas():
    conn = conectar()
    c = conn.cursor()
    c.executescript("""
        CREATE TABLE IF NOT EXISTS plantonistas (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            nome TEXT NOT NULL,
            matricula TEXT NOT NULL,
            cpf TEXT,
            telefone TEXT
        );
        CREATE TABLE IF NOT EXISTS escalas (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            data_inicio TEXT NOT NULL,
            data_fim TEXT NOT NULL,
            turno TEXT NOT NULL,
            vagas INTEGER NOT NULL,
            plantonistas TEXT NOT NULL,
            viatura_id INTEGER,
            coordenador_id INTEGER,
            FOREIGN KEY (viatura_id) REFERENCES viaturas(id),
            FOREIGN KEY (coordenador_id) REFERENCES coordenadores(id)
        );
        CREATE TABLE IF NOT EXISTS historico (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            data_inicio TEXT NOT NULL,
            data_fim TEXT NOT NULL,
            turno TEXT NOT NULL,
            plantonistas TEXT NOT NULL,
            horas_normais REAL,
            horas_especiais REAL
        );
        CREATE TABLE IF NOT EXISTS viaturas (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            placa TEXT NOT NULL,
            modelo TEXT
        );
        CREATE TABLE IF NOT EXISTS coordenadores (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            nome TEXT NOT NULL,
            matricula TEXT,
            contato TEXT
        );
    """)
    conn.commit()
    conn.close()


# --- Plantonistas ---
def listar_plantonistas():
    with conectar() as conn:
        return pd.read_sql_query("SELECT * FROM plantonistas", conn)

def cadastrar_plantonista(nome, matricula, cpf, telefone):
    with conectar() as conn:
        conn.execute("INSERT INTO plantonistas (nome, matricula, cpf, telefone) VALUES (?, ?, ?, ?)", (nome, matricula, cpf, telefone))

def apagar_plantonista(id_plantonista):
    with conectar() as conn:
        conn.execute("DELETE FROM plantonistas WHERE id=?", (id_plantonista,))

# --- Escalas ---

def gerar_escala_manual(data_inicio, data_fim, turno, vagas, plantonistas, viatura_id, coordenador_id):
    plantonistas_str = json.dumps(plantonistas, ensure_ascii=False)
    horas_normais, horas_especiais = calcular_horas_extras(data_inicio, data_fim)

    with conectar() as conn:
        conn.execute(
            """
            INSERT INTO escalas (data_inicio, data_fim, turno, vagas, plantonistas, viatura_id, coordenador_id)
            VALUES (?, ?, ?, ?, ?, ?, ?)
            """,
            (data_inicio, data_fim, turno, vagas, plantonistas_str, viatura_id, coordenador_id)
        )
        conn.execute(
            """
            INSERT INTO historico (data_inicio, data_fim, turno, plantonistas, horas_normais, horas_especiais)
            VALUES (?, ?, ?, ?, ?, ?)
            """,
            (data_inicio, data_fim, turno, plantonistas_str, horas_normais, horas_especiais)
        )


def gerar_escala_automatica(data_inicio, data_fim, turno, vagas):
    df = listar_plantonistas()
    selecionados = df['nome'].tolist()[:vagas]
    gerar_escala_manual(data_inicio, data_fim, turno, vagas, selecionados)

def apagar_escala(id_escala):
    with conectar() as conn:
        c = conn.cursor()
        c.execute("SELECT data_inicio, data_fim, turno FROM escalas WHERE id=?", (id_escala,))
        escala = c.fetchone()
        if escala:
            data_inicio, data_fim, turno = escala
            c.execute("DELETE FROM escalas WHERE id=?", (id_escala,))
            c.execute("DELETE FROM historico WHERE data_inicio=? AND data_fim=? AND turno=?", (data_inicio, data_fim, turno))

# --- Viaturas ---
def listar_viaturas():
    with conectar() as conn:
        return pd.read_sql_query("SELECT * FROM viaturas", conn)

def cadastrar_viatura(placa, modelo):
    with conectar() as conn:
        conn.execute("INSERT INTO viaturas (placa, modelo) VALUES (?, ?)", (placa, modelo))

def apagar_viatura(id_viatura):
    with conectar() as conn:
        conn.execute("DELETE FROM viaturas WHERE id=?", (id_viatura,))

# --- Coordenadores ---
def listar_coordenadores():
    with conectar() as conn:
        return pd.read_sql_query("SELECT * FROM coordenadores", conn)

def cadastrar_coordenador(nome, matricula, contato):
    with conectar() as conn:
        conn.execute("INSERT INTO coordenadores (nome, matricula, contato) VALUES (?, ?, ?)", (nome, matricula, contato))

def apagar_coordenador(id_coordenador):
    with conectar() as conn:
        conn.execute("DELETE FROM coordenadores WHERE id=?", (id_coordenador,))

# --- Histórico ---
def gerar_historico_excel_por_equipe():
    with conectar() as conn:
        df = pd.read_sql_query("SELECT * FROM escalas", conn)
    df['equipe'] = df['plantonistas'].apply(safe_json_loads)
    df[['data_inicio', 'data_fim', 'turno', 'vagas', 'equipe']].to_excel("relatorios/historico_por_equipe.xlsx", index=False)

def gerar_historico_pdf_por_equipe():
    from fpdf import FPDF
    with conectar() as conn:
        df = pd.read_sql_query("SELECT * FROM escalas", conn)
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Arial", size=12)
    pdf.cell(200, 10, txt="Histórico de Escalas por Equipe", ln=True, align='C')
    for _, row in df.iterrows():
        equipe = safe_json_loads(row['plantonistas'])
        pdf.cell(0, 10, f"{row['data_inicio']} - {row['data_fim']} | {row['turno']} | Vagas: {row['vagas']} | {equipe}", ln=True)
    os.makedirs("relatorios", exist_ok=True)
    pdf.output("relatorios/historico_por_equipe.pdf")

# --- Utilitários ---
def calcular_horas_extras(data_inicio, data_fim):
    di = datetime.strptime(data_inicio, '%Y-%m-%d %H:%M')
    df = datetime.strptime(data_fim, '%Y-%m-%d %H:%M')
    atual = di
    horas_normais = horas_especiais = 0
    while atual < df:
        if atual.weekday() >= 5:
            horas_especiais += 1 / 60
        else:
            if 6 <= atual.hour < 24:
                horas_normais += 1 / 60
            else:
                horas_especiais += 1 / 60
        atual += timedelta(minutes=1)
    return round(horas_normais, 2), round(horas_especiais, 2)

def safe_json_loads(x):
    try:
        if isinstance(x, str) and x.strip().startswith("["):
            return ", ".join(json.loads(x))
        return x or ''
    except Exception:
        return x or ''

def safe_list_load(x):
    try:
        return json.loads(x) if isinstance(x, str) and x.strip().startswith("[") else [x] if x else []
    except Exception:
        return [x] if x else []


def docx_para_pdf(docx_path, pdf_dir):
    sistema = platform.system()
    
    executable = "soffice" if sistema == "Windows" else "libreoffice"

    try:
        subprocess.run([
            executable,
            "--headless",
            "--convert-to", "pdf",
            "--outdir", pdf_dir,
            docx_path
        ], check=True)

        # Renomeia o PDF gerado para um nome padrão
        generated_pdf = os.path.join(pdf_dir, os.path.basename(docx_path).replace(".docx", ".pdf"))
        final_pdf_path = os.path.join(pdf_dir, "escala_completa.pdf")
        os.rename(generated_pdf, final_pdf_path)
        return final_pdf_path

    except FileNotFoundError:
        raise FileNotFoundError(f"'{executable}' não foi encontrado. Instale o LibreOffice e adicione ao PATH.")
    except subprocess.CalledProcessError as e:
        raise RuntimeError(f"Erro na conversão do DOCX para PDF: {e}")
    
dias_semana = {
"Monday": "SEGUNDA-FEIRA",
"Tuesday": "TERÇA-FEIRA",
"Wednesday": "QUARTA-FEIRA",
"Thursday": "QUINTA-FEIRA",
"Friday": "SEXTA-FEIRA",
"Saturday": "SÁBADO",
"Sunday": "DOMINGO"
}


meses_pt = {
    1: "janeiro", 2: "fevereiro", 3: "março", 4: "abril",
    5: "maio", 6: "junho", 7: "julho", 8: "agosto",
    9: "setembro", 10: "outubro", 11: "novembro", 12: "dezembro"
}

hoje = datetime.now()
data_hoje = f"{hoje.day} de {meses_pt[hoje.month]} de {hoje.year}"



def gerar_pdf_escala_por_equipe(ids=None):
    conn = conectar()
    query = "SELECT * FROM escalas"
    if ids:
        placeholders = ','.join(['?'] * len(ids))
        query += f" WHERE id IN ({placeholders})"
        df = pd.read_sql_query(query, conn, params=ids)
    else:
        df = pd.read_sql_query(query, conn)

    df['plantonistas'] = df['plantonistas'].apply(safe_list_load)
    plantonistas_db = pd.read_sql_query('SELECT * FROM plantonistas', conn)

    os.makedirs("relatorios", exist_ok=True)

    hoje = datetime.now()
    data_hoje = f"{hoje.day} de {meses_pt[hoje.month]} de {hoje.year}"
  # Ex: 19 de maio de 2025

    with tempfile.TemporaryDirectory() as tmpdir:
        docx_paths = []

        for idx, row in df.iterrows():
            doc = Document("base_escala.docx")

            data_inicio = datetime.strptime(row['data_inicio'], '%Y-%m-%d %H:%M')
            data_fim = datetime.strptime(row['data_fim'], '%Y-%m-%d %H:%M')

            dia_semana = dias_semana[data_inicio.strftime('%A')]
            dia_semana_fim = dias_semana[data_fim.strftime('%A')]

            data_formatada = data_inicio.strftime('%d/%m/%Y')
            data_fim_formatada = data_fim.strftime('%d/%m/%Y')
            turno = row['turno']
            total = len(row['plantonistas'])

            placa = '---'
            if row['viatura_id']:
                viatura_row = conn.execute("SELECT placa FROM viaturas WHERE id = ?", (row['viatura_id'],)).fetchone()
                placa = viatura_row['placa'] if viatura_row else '---'

            coordenador = '---'
            if row['coordenador_id']:
                coord_row = conn.execute("SELECT nome FROM coordenadores WHERE id = ?", (row['coordenador_id'],)).fetchone()
                coordenador = coord_row['nome'] if coord_row else '---'

            for p in doc.paragraphs:
                p.text = p.text.replace("{{dia_semana}}", dia_semana)
                p.text = p.text.replace("{{data}}", data_formatada)
                p.text = p.text.replace("{{dia_semana_fim}}", dia_semana_fim)
                p.text = p.text.replace("{{data_fim}}", data_fim_formatada)
                p.text = p.text.replace("{{turno}}", turno)
                p.text = p.text.replace("{{placa}}", placa)
                p.text = p.text.replace("{{total}}", str(total))
                p.text = p.text.replace("{{coordenador}}", coordenador)
                p.text = p.text.replace("{{data_hoje}}", data_hoje)

            tabela = next((t for t in doc.tables if "Matrícula" in t.cell(0, 1).text), None)

            if tabela:
                for nome in row['plantonistas']:
                    match = plantonistas_db[plantonistas_db['nome'] == nome]
                    dados = match.iloc[0] if not match.empty else {}

                    linha = tabela.add_row().cells
                    linha[0].paragraphs[0].add_run(f"OIP {nome}")
                    linha[1].paragraphs[0].add_run(dados.get('matricula', '---'))
                    linha[2].paragraphs[0].add_run(dados.get('cpf', '---'))
                    if len(linha) > 3:
                        linha[3].paragraphs[0].add_run(dados.get('telefone', '---'))

            assinatura_path = "assinatura.png"
            if os.path.exists(assinatura_path):
                doc.add_paragraph("")
                doc.add_picture(assinatura_path, width=Inches(2.5))
                doc.add_paragraph("DR MARCOS VINÍCIUS CACAU DE LIMA\nDelegado De Polícia Civil")

            temp_docx = os.path.join(tmpdir, f"escala_{row['id']}.docx")
            doc.save(temp_docx)
            docx_paths.append(temp_docx)

        final_docx_path = os.path.join("relatorios", "escala_completa.docx")
        merged = Document(docx_paths[0])
        for other_path in docx_paths[1:]:
            sub_doc = Document(other_path)
            for element in sub_doc.element.body:
                merged.element.body.append(element)
        merged.save(final_docx_path)

        final_pdf_path = docx_para_pdf(final_docx_path, "relatorios")

        with open(final_pdf_path, "rb") as f:
            return f.read()
